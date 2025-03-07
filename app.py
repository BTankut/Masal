import os
import base64
import io
import sys
import traceback
import json
import logging
import time
import datetime
import glob
import argparse
from logging.handlers import RotatingFileHandler
from flask import Flask, render_template, request, jsonify, send_file, make_response, after_this_request
from flask_cors import CORS
from dotenv import load_dotenv
import google.generativeai as genai
from PIL import Image
import requests
from openai import OpenAI
import openai  # RateLimitError gibi hata tiplerini yakalamak için
import tempfile
from gtts import gTTS
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.shared import Inches

# Log klasörünü oluştur
logs_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'logs')
os.makedirs(logs_dir, exist_ok=True)

# Geçerli oturum için benzersiz bir log dosya ismi oluştur
timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
app_log_path = os.path.join(logs_dir, f"app_{timestamp}.log")
debug_log_path = os.path.join(logs_dir, f"debug_{timestamp}.log")
prompt_log_path = os.path.join(logs_dir, f"prompt_{timestamp}.log")

# Loglama yapılandırması
logger = logging.getLogger("masal_app")
logger.setLevel(logging.DEBUG)

# Log formatı
log_format = logging.Formatter('%(asctime)s [%(levelname)s] %(message)s')

# Console handler
console_handler = logging.StreamHandler(sys.stdout)
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(log_format)

# Normal log dosyası (INFO ve üstü seviye mesajlar için)
file_handler = RotatingFileHandler(
    app_log_path, maxBytes=10*1024*1024, backupCount=3
)
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(log_format)

# Debug log dosyası (tüm detaylı loglar için)
debug_file_handler = RotatingFileHandler(
    debug_log_path, maxBytes=10*1024*1024, backupCount=3
)
debug_file_handler.setLevel(logging.DEBUG)
debug_file_handler.setFormatter(log_format)

# Handlerleri loggera ekle
logger.addHandler(console_handler)
logger.addHandler(file_handler)
logger.addHandler(debug_file_handler)

# Promt testi için ayrı bir logger
prompt_logger = logging.getLogger("prompt_test")
prompt_logger.setLevel(logging.INFO)
prompt_handler = RotatingFileHandler(
    prompt_log_path, maxBytes=10*1024*1024, backupCount=3
)
prompt_handler.setFormatter(log_format)
prompt_logger.addHandler(prompt_handler)
prompt_logger.addHandler(console_handler)

# .env dosyasından API anahtarlarını yükle
load_dotenv()
google_api_key = os.getenv("GOOGLE_API_KEY")
openai_api_key = os.getenv("OPENAI_API_KEY")

if not google_api_key:
    logger.error("Google API anahtarı bulunamadı! .env dosyanızı kontrol edin.")
    raise ValueError("Lütfen .env dosyasına GOOGLE_API_KEY değerini ekleyin.")

# Gemini API'yi yapılandır
genai.configure(api_key=google_api_key)
logger.info("Gemini API yapılandırıldı.")

# OpenAI API'yi yapılandır
if openai_api_key:
    openai_client = OpenAI(api_key=openai_api_key)
    logger.info("OpenAI API yapılandırıldı.")
else:
    openai_client = None
    logger.warning("UYARI: OpenAI API anahtarı bulunamadı. DALL-E görsel oluşturma devre dışı.")

app = Flask(__name__)
# Tüm kaynaklardan gelen isteklere izin ver
CORS(app, resources={r"/*": {"origins": "*"}})
logger.info("Flask uygulaması ve CORS yapılandırıldı.")

@app.route('/')
def index():
    logger.debug("Ana sayfa isteği alındı.")
    response = make_response(render_template('index.html'))
    # Önbelleği devre dışı bırak
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response
    
@app.route('/debug')
def debug_page():
    logger.debug("Debug sayfası isteği alındı.")
    response = make_response(render_template('debug.html'))
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

def parse_form_data(request):
    """Form veya JSON verilerini işler ve karakter özelliklerini içeren sözlük döndürür"""
    # Form verilerini al
    if request.is_json:
        data = request.json
        character_name = data.get('character_name', '')
        character_type = data.get('character_type', '')
        setting = data.get('setting', '')
        theme = data.get('theme', '')
        
        try:
            word_limit = int(data.get('word_limit', 200))
        except ValueError:
            logger.warning("Kelime sayısı geçersiz, varsayılan değer kullanılıyor.")
            word_limit = 200
            
        image_api = data.get('image_api', 'dalle')
        text_api = data.get('text_api', 'openai')
        
        # Karakter detaylarını al
        character_attributes = {
            'character_age': data.get('character_age', ''),
            'character_gender': data.get('character_gender', ''),
            'character_hair_color': data.get('character_hair_color', ''),
            'character_hair_type': data.get('character_hair_type', ''),
            'character_skin_color': data.get('character_skin_color', '')
        }
    else:
        # Form verilerini al
        character_name = request.form.get('character_name', '')
        character_type = request.form.get('character_type', '')
        setting = request.form.get('setting', '')
        theme = request.form.get('theme', '')
        
        # word_limit ve api değerleri için hata kontrolü
        try:
            word_limit = int(request.form.get('word_limit', 200))
        except ValueError:
            logger.warning("Kelime sayısı geçersiz, varsayılan değer kullanılıyor.")
            word_limit = 200
            
        image_api = request.form.get('image_api', 'dalle')
        text_api = request.form.get('text_api', 'openai')
        
        # Karakter detaylarını al
        character_attributes = {
            'character_age': request.form.get('character_age', ''),
            'character_gender': request.form.get('character_gender', ''),
            'character_hair_color': request.form.get('character_hair_color', ''),
            'character_hair_type': request.form.get('character_hair_type', ''),
            'character_skin_color': request.form.get('character_skin_color', '')
        }
        
    # Karakter özelliklerini logla
    for key, value in character_attributes.items():
        logger.info(f"Karakter özellik: {key} = {value}")
        
    return character_name, character_type, setting, theme, word_limit, image_api, text_api, character_attributes

@app.route('/generate_tale', methods=['POST'])
def generate_tale():
    try:
        logger.info("Masal oluşturma isteği alındı.")
        
        # Request içeriğini logla
        try:
            logger.debug(f"İstek türü: {request.content_type}")
            if request.data:
                logger.debug(f"İstek verisi: {request.data.decode('utf-8')}")
            logger.debug(f"Form verisi: {request.form}")
            logger.debug(f"JSON verisi: {request.json if request.is_json else 'JSON verisi yok'}")
        except Exception as e:
            logger.error(f"İstek verisi loglanırken hata: {str(e)}")
        
        # Form verilerini parse et
        character_name, character_type, setting, theme, word_limit, image_api, text_api, character_attributes = parse_form_data(request)
        
        logger.info(f"Alınan form verileri: Karakter Adı: {character_name}, Karakter Türü: {character_type}, Karakter Özellikleri: {character_attributes}, Ortam: {setting}, Tema: {theme}, Kelime Sayısı: {word_limit}, Görsel API: {image_api}, Metin API: {text_api}")
        
        # Veri doğrulama
        if not character_name or not character_type or not setting or not theme:
            logger.error("Eksik form verileri")
            return jsonify({"error": "Lütfen tüm alanları doldurunuz."}), 400
        
        # Masal oluşturma
        logger.info("Masal metni oluşturuluyor...")
        tale_text = generate_tale_text(character_name, character_type, setting, theme, word_limit, text_api, character_attributes)
        
        # Başlık oluştur
        tale_title = f"{character_name}'nin {setting} Macerası"
        logger.info(f"Masal başlığı oluşturuldu: {tale_title}")
        
        # İlk sayfa için görsel oluştur
        logger.info(f"{image_api} API kullanarak ilk sayfa görseli oluşturuluyor...")
        
        # Masalı bölümlere ayır
        sections = split_text_into_sections(tale_text, 50)
        if sections:
            first_section = sections[0]
            # Karakter özelliklerini al ve prompta ekle
            character_description = ""
            # Debug log için özellikleri göster
            logger.info(f"character_age: {request.form.get('character_age', 'YOK')}")
            logger.info(f"character_gender: {request.form.get('character_gender', 'YOK')}")
            logger.info(f"character_hair_color: {request.form.get('character_hair_color', 'YOK')}")
            logger.info(f"character_hair_type: {request.form.get('character_hair_type', 'YOK')}")
            logger.info(f"character_skin_color: {request.form.get('character_skin_color', 'YOK')}")
            
            if request.form.get('character_age'):
                character_description += f"{request.form.get('character_age')} yaşında, "
                
            if request.form.get('character_gender'):
                character_description += f"{request.form.get('character_gender')}, "
            
            hair_parts = []
            if request.form.get('character_hair_type'):
                hair_parts.append(request.form.get('character_hair_type'))
            if request.form.get('character_hair_color'):
                hair_parts.append(request.form.get('character_hair_color'))
            
            if hair_parts:
                character_description += f"{' '.join(hair_parts)} saçlı, "
            
            if request.form.get('character_skin_color'):
                character_description += f"{request.form.get('character_skin_color')} tenli, "
            
            # Son virgülü ve boşluğu kaldır
            if character_description:
                character_description = character_description.rstrip(", ")
                character_description = f"({character_description}) "
            
            # İlk sayfa için özel prompt oluştur
            image_prompt = f"{character_name} adlı {character_description}{character_type} karakteri {setting} ortamında: {first_section}"
            # Prompt'u logla
            prompt_logger.info(f"Image prompt: {image_prompt}")
        else:
            # Bölüm yoksa genel bir prompt kullan
            character_description = ""
            # Debug log için özellikleri göster
            logger.info(f"Bölüm yoksa - character_age: {request.form.get('character_age', 'YOK')}")
            logger.info(f"Bölüm yoksa - character_gender: {request.form.get('character_gender', 'YOK')}")
            logger.info(f"Bölüm yoksa - character_hair_color: {request.form.get('character_hair_color', 'YOK')}")
            logger.info(f"Bölüm yoksa - character_hair_type: {request.form.get('character_hair_type', 'YOK')}")
            logger.info(f"Bölüm yoksa - character_skin_color: {request.form.get('character_skin_color', 'YOK')}")
            
            if request.form.get('character_age'):
                character_description += f"{request.form.get('character_age')} yaşında, "
                
            if request.form.get('character_gender'):
                character_description += f"{request.form.get('character_gender')}, "
            
            hair_parts = []
            if request.form.get('character_hair_type'):
                hair_parts.append(request.form.get('character_hair_type'))
            if request.form.get('character_hair_color'):
                hair_parts.append(request.form.get('character_hair_color'))
            
            if hair_parts:
                character_description += f"{' '.join(hair_parts)} saçlı, "
            
            if request.form.get('character_skin_color'):
                character_description += f"{request.form.get('character_skin_color')} tenli, "
            
            # Son virgülü ve boşluğu kaldır
            if character_description:
                character_description = character_description.rstrip(", ")
                character_description = f"({character_description}) "
                
            image_prompt = f"{character_name} adlı {character_description}{character_type} karakteri {setting} ortamında, {theme} temalı bir masal için illüstrasyon"
            # Prompt'u logla
            prompt_logger.info(f"Image prompt: {image_prompt}")
        
        # Görseli oluştur
        image_data = generate_image_for_section(image_prompt, image_api)
        
        # Görsel oluşturma başarı/başarısızlık durumunu logla
        if image_data:
            logger.info("Görsel başarıyla oluşturuldu.")
        else:
            logger.warning("Görsel oluşturma başarısız, varsayılan görsel kullanılacak.")
        
        # Yanıt hazırla
        response_data = {
            "tale_title": tale_title,
            "tale_text": tale_text,
            "image_url": f"data:image/jpeg;base64,{image_data}" if image_data else "static/img/default-tale.jpg"
        }
        
        logger.info("Masal başarıyla oluşturuldu.")
        return jsonify(response_data)
        
    except Exception as e:
        error_details = {
            "error": str(e),
            "traceback": traceback.format_exc()
        }
        logger.error(f"Masal oluşturulurken hata oluştu: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({"error": str(e), "details": error_details}), 500

@app.route('/save_word', methods=['POST'])
def save_word():
    try:
        data = request.json
        tale_text = data.get('tale_text', '')
        images = data.get('images', [])
        
        # İstek değerlerini logla
        logger.info(f"Word dosyası oluşturma isteği alındı - Metin uzunluğu: {len(tale_text)} karakter, Görsel sayısı: {len(images)}")
        
        # Word dosyası oluştur
        doc_path = create_word_document(tale_text, images)
        
        logger.info(f"Word dosyası başarıyla oluşturuldu: {doc_path}")
        
        # Word dosyasını gönder
        return send_file(doc_path, as_attachment=True, download_name='masal.docx')
    except Exception as e:
        logger.error(f"Word dosyası oluşturma hatası: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/generate_audio', methods=['POST'])
def generate_audio():
    try:
        data = request.json
        text = data.get('text', '')
        page = data.get('page', 0)  # Sayfa numarasını al (debug için)
        
        logger.info(f"Sayfa {page+1} için ses oluşturma isteği alındı, metin uzunluğu: {len(text)} karakter")
        logger.debug(f"Sayfa {page+1} metin başlangıcı: {text[:30]}...")
        
        # Ses dosyası oluştur
        audio_path = create_audio(text)
        
        logger.info(f"Sayfa {page+1} için ses dosyası oluşturuldu: {audio_path}")
        
        # Ses dosyasını gönder - mimetype belirtiyoruz ve attachment olarak değil normal dosya olarak gönderiyoruz
        return send_file(audio_path, mimetype='audio/mpeg', as_attachment=False)
    except Exception as e:
        logger.error(f"Sayfa {page+1} için ses oluşturma hatası: {str(e)}")
        
@app.route('/save_tale', methods=['POST'])
def save_tale():
    try:
        data = request.json
        if not data:
            logger.error("save_tale: Gelen veri boş")
            return jsonify({'error': 'Veri bulunamadı'}), 400
        
        # Masal ID'sini kontrol et, yoksa oluştur
        tale_id = data.get('id', str(int(time.time())))
        tale_type = data.get('type', 'history')  # 'history' veya 'favorites'
        tale_title = data.get('title', 'Başlıksız Masal')
        
        logger.info(f"save_tale: Masal kaydediliyor - ID: {tale_id}, Tür: {tale_type}, Başlık: {tale_title}")
        
        # Tek bir dizin kullan
        directory = 'static/tales/all'
        
        # Klasör yoksa oluştur
        os.makedirs(directory, exist_ok=True)
        
        # Masal bilgilerini JSON olarak kaydet
        tale_info = {
            'id': tale_id,
            'title': tale_title,
            'text': data.get('text', ''),
            'date': data.get('date', datetime.datetime.now().isoformat()),
            'characterName': data.get('characterName', ''),
            'characterType': data.get('characterType', ''),
            'setting': data.get('setting', ''),
            'theme': data.get('theme', ''),
            'type': tale_type,  # Masal tipini (history veya favorites) JSON içine de ekle
            'isFavorite': tale_type == 'favorites'  # Favorilere ait mi işareti ekle
        }
        
        # JSON dosyasını kaydet
        json_path = f"{directory}/{tale_id}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(tale_info, f, ensure_ascii=False, indent=2)
        
        logger.info(f"save_tale: JSON bilgileri kaydedildi: {json_path}")
        
        # Görsel varsa kaydet
        image_saved = False
        if 'image' in data and data['image']:
            # Base64 veriyi ayır ve decode et
            try:
                image_path = f"{directory}/{tale_id}_image.jpg"
                if data['image'].startswith('data:image/'):
                    logger.debug(f"save_tale: Base64 görsel verisi işleniyor")
                    img_data = data['image'].split(',')[1]
                    binary_data = base64.b64decode(img_data)
                else:
                    # URL ise indirip kaydedelim
                    logger.debug(f"save_tale: URL'den görsel indiriliyor: {data['image'][:30]}...")
                    response = requests.get(data['image'])
                    binary_data = response.content
                
                # Görsel dosyasını kaydet
                with open(image_path, 'wb') as f:
                    f.write(binary_data)
                
                logger.info(f"save_tale: Ana görsel kaydedildi: {image_path}")
                image_saved = True
            except Exception as img_err:
                logger.error(f"save_tale: Ana görsel kaydedilirken hata: {img_err}")
        
        if not image_saved:
            logger.warning(f"save_tale: Masal görseli bulunamadı veya kaydedilemedi")
        
        # Eğer sayfalanmış metin varsa her sayfa için ayrı resim kaydedelim
        pages_count = 0
        pages_with_images = 0
        
        if 'pages' in data and isinstance(data['pages'], list):
            pages_count = len(data['pages'])
            logger.info(f"save_tale: {pages_count} sayfa verisi kaydedilecek")
            
            for i, page in enumerate(data['pages']):
                page_content = page.get('text', '')
                page_image = page.get('image', '')
                
                # Sayfa metnini kaydet
                page_text_path = f"{directory}/{tale_id}_page_{i}.txt"
                with open(page_text_path, 'w', encoding='utf-8') as f:
                    f.write(page_content)
                
                # Sayfa görselini kaydet
                if page_image:
                    try:
                        page_image_path = f"{directory}/{tale_id}_page_{i}_image.jpg"
                        
                        if page_image.startswith('data:image/'):
                            img_data = page_image.split(',')[1]
                            binary_data = base64.b64decode(img_data)
                        else:
                            # URL ise indirip kaydedelim
                            response = requests.get(page_image)
                            binary_data = response.content
                        
                        with open(page_image_path, 'wb') as f:
                            f.write(binary_data)
                        
                        pages_with_images += 1
                    except Exception as page_img_err:
                        logger.error(f"save_tale: Sayfa {i+1} görseli kaydedilirken hata: {page_img_err}")
        
        logger.info(f"save_tale: Toplam {pages_count} sayfa kaydedildi, {pages_with_images} sayfa görsel içeriyor")
        
        # Eğer ses dosyaları varsa kaydedelim
        audio_count = 0
        if 'audios' in data and isinstance(data['audios'], dict):
            for page_idx, audio_data in data['audios'].items():
                if 'blob' in audio_data:
                    try:
                        # Base64 veri
                        audio_path = f"{directory}/{tale_id}_page_{page_idx}_audio.mp3"
                        binary_data = base64.b64decode(audio_data['blob'])
                        with open(audio_path, 'wb') as f:
                            f.write(binary_data)
                        audio_count += 1
                    except Exception as audio_err:
                        logger.error(f"save_tale: Sayfa {page_idx} ses dosyası kaydedilirken hata: {audio_err}")
        
        logger.info(f"save_tale: Toplam {audio_count} sayfa için ses dosyası kaydedildi")
        logger.info(f"save_tale: Masal başarıyla kaydedildi - ID: {tale_id}, Tür: {tale_type}, Başlık: {tale_title}")
        
        return jsonify({
            'success': True,
            'message': 'Masal başarıyla kaydedildi',
            'tale_id': tale_id
        })
    
    except Exception as e:
        logger.error(f"Masal kaydetme hatası: {e}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/list_tales', methods=['GET'])
def list_tales():
    try:
        tale_type = request.args.get('type', 'history')  # 'history' veya 'favorites'
        
        logger.info(f"list_tales: {tale_type} tipindeki masallar listeleniyor")
        
        # Tek klasör kullan
        directory = 'static/tales/all'
        
        # Klasör yoksa oluştur
        os.makedirs(directory, exist_ok=True)
        
        # JSON dosyalarını bul
        tales = []
        json_files = glob.glob(f"{directory}/*.json")
        
        logger.info(f"list_tales: {len(json_files)} adet JSON dosyası bulundu")
        
        for json_file in json_files:
            try:
                with open(json_file, 'r', encoding='utf-8') as f:
                    tale_data = json.load(f)
                
                # Dosya adından ID al
                tale_id = os.path.basename(json_file).replace('.json', '')
                
                # Masal tipine göre filtrele - JSON içindeki type değerine göre
                tale_data_type = tale_data.get('type', 'history')
                if tale_type != 'all' and tale_data_type != tale_type:
                    # İstenen tipe uymayan masalları atla
                    continue
                
                # Favori ise özellikle belirt
                is_favorite = tale_data.get('isFavorite', False) or tale_data_type == 'favorites'
                
                # Görsel ve ses dosyalarını kontrol et
                image_path = f"{directory}/{tale_id}_image.jpg"
                has_image = os.path.exists(image_path)
                
                # Sayfa dosyalarını say
                page_files = glob.glob(f"{directory}/{tale_id}_page_*.txt")
                audio_files = glob.glob(f"{directory}/{tale_id}_page_*_audio.mp3")
                page_image_files = glob.glob(f"{directory}/{tale_id}_page_*_image.jpg")
                
                # Masal bilgilerini listeye ekle
                tale_info = {
                    'id': tale_id,
                    'title': tale_data.get('title', 'Başlıksız Masal'),
                    'date': tale_data.get('date', ''),
                    'characterName': tale_data.get('characterName', ''),
                    'characterType': tale_data.get('characterType', ''),
                    'has_image': has_image,
                    'image_url': image_path if has_image else None,
                    'page_count': len(page_files),
                    'has_audio': len(audio_files) > 0,
                    'has_page_images': len(page_image_files) > 0,
                    'type': tale_data_type,
                    'isFavorite': is_favorite
                }
                
                logger.debug(f"list_tales: Masal bilgileri: {tale_info}")
                tales.append(tale_info)
                
            except Exception as e:
                logger.error(f"list_tales: JSON dosyası okuma hatası ({json_file}): {e}")
        
        # Tarihe göre sırala (en yeni en üstte)
        tales.sort(key=lambda x: x.get('date', ''), reverse=True)
        
        # En fazla 5 masal döndür
        tales = tales[:5]
        
        logger.info(f"list_tales: Toplam {len(tales)} adet masal listelendi")
        
        return jsonify(tales)
    
    except Exception as e:
        logger.error(f"Masalları listeleme hatası: {e}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/clear_tales', methods=['POST'])
def clear_tales():
    try:
        data = request.json
        tale_type = data.get('type', 'all')  # 'history', 'favorites' veya 'all'
        
        logger.info(f"clear_tales: {tale_type} tipindeki masallar temizleniyor")
        
        directory = 'static/tales/all'
        
        # Tek klasör olduğu için, favoriler veya geçmişe göre filtreleyerek silme yapmalıyız
        if os.path.exists(directory):
            if tale_type == 'all':
                # Tüm dosyaları temizle
                try:
                    for file in os.listdir(directory):
                        file_path = os.path.join(directory, file)
                        if os.path.isfile(file_path):
                            os.unlink(file_path)
                    logger.info(f"clear_tales: Tüm masallar silindi")
                except Exception as e:
                    logger.error(f"clear_tales: Masallar silinirken hata: {e}")
            else:
                # Belirli tipteki masalları bul ve sil
                json_files = glob.glob(f"{directory}/*.json")
                for json_file in json_files:
                    try:
                        # JSON dosyasını oku ve tipini kontrol et
                        with open(json_file, 'r', encoding='utf-8') as f:
                            tale_data = json.load(f)
                        
                        tale_data_type = tale_data.get('type', 'history')
                        if tale_data_type == tale_type:
                            tale_id = os.path.basename(json_file).replace('.json', '')
                            
                            # Bu masala ait tüm dosyaları sil
                            for file in os.listdir(directory):
                                if file.startswith(tale_id):
                                    file_path = os.path.join(directory, file)
                                    if os.path.isfile(file_path):
                                        os.unlink(file_path)
                            
                            logger.info(f"clear_tales: {tale_id} ID'li {tale_type} masalı silindi")
                    except Exception as e:
                        logger.error(f"clear_tales: Masal silinirken hata: {e}")
        
        return jsonify({'success': True, 'message': 'Masallar başarıyla temizlendi'})
    
    except Exception as e:
        logger.error(f"Masalları temizleme hatası: {e}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/load_tale/<tale_id>')
def load_tale(tale_id):
    try:
        tale_type = request.args.get('type', 'history')  # 'history' veya 'favorites' (artık sadece bilgi amaçlı)
        
        logger.info(f"load_tale: Masal {tale_id} yükleniyor (istenen tip: {tale_type})")
        
        # Tek klasör kullan
        directory = 'static/tales/all'
        
        # JSON dosyasını oku
        json_file = f"{directory}/{tale_id}.json"
        
        if not os.path.exists(json_file):
            logger.error(f"load_tale: Masal bulunamadı - dosya yok: {json_file}")
            return jsonify({'error': 'Masal bulunamadı'}), 404
        
        with open(json_file, 'r', encoding='utf-8') as f:
            tale_data = json.load(f)
        
        logger.info(f"load_tale: Masal JSON verisi yüklendi: {tale_data.get('title', 'Başlıksız')}")
        
        # Tale data içinde zaten type var, olmasa bile alalım
        data_type = tale_data.get('type', tale_type) 
        tale_data['type'] = data_type
        
        # Favori durumunu belirt
        is_favorite = tale_data.get('isFavorite', False) or data_type == 'favorites'
        tale_data['isFavorite'] = is_favorite
        
        # Masala ait görselleri bul
        tale_data['pages'] = []
        
        # Ana görsel
        image_path = f"{directory}/{tale_id}_image.jpg"
        if os.path.exists(image_path):
            tale_data['image_url'] = image_path
            logger.info(f"load_tale: Ana görsel bulundu: {image_path}")
        else:
            logger.warning(f"load_tale: Ana görsel bulunamadı: {image_path}")
        
        # Sayfa dosyalarını bul
        page_files = sorted(glob.glob(f"{directory}/{tale_id}_page_*.txt"))
        logger.info(f"load_tale: {len(page_files)} adet sayfa metni dosyası bulundu")
        
        for i, page_file in enumerate(page_files):
            try:
                page_idx = os.path.basename(page_file).replace(f'{tale_id}_page_', '').replace('.txt', '')
                
                # Sayfa metni
                with open(page_file, 'r', encoding='utf-8') as f:
                    page_text = f.read()
                
                # Sayfa görseli
                page_image = f"{directory}/{tale_id}_page_{page_idx}_image.jpg"
                page_image_url = None
                if os.path.exists(page_image):
                    page_image_url = page_image
                    logger.debug(f"load_tale: Sayfa {page_idx} için görsel bulundu")
                
                # Sayfa ses dosyası
                page_audio = f"{directory}/{tale_id}_page_{page_idx}_audio.mp3"
                page_audio_url = None
                if os.path.exists(page_audio):
                    page_audio_url = page_audio
                    logger.debug(f"load_tale: Sayfa {page_idx} için ses dosyası bulundu")
                
                # Sayfa bilgilerini ekle
                tale_data['pages'].append({
                    'index': int(page_idx),
                    'text': page_text,
                    'image_url': page_image_url,
                    'audio_url': page_audio_url
                })
            except Exception as page_err:
                logger.error(f"load_tale: Sayfa {i} yüklenirken hata: {page_err}")
        
        # Sayfaları sırala
        tale_data['pages'].sort(key=lambda x: x['index'])
        
        logger.info(f"load_tale: Masal başarıyla yüklendi - {tale_data.get('title', 'Başlıksız')} ({len(tale_data['pages'])} sayfa)")
        
        return jsonify(tale_data)
    
    except Exception as e:
        logger.error(f"Masal yükleme hatası: {e}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

@app.route('/generate_page_image', methods=['POST'])
def generate_page_image():
    try:
        data = request.json
        page_text = data.get('page_text', '')
        character_name = data.get('character_name', '')
        character_type = data.get('character_type', '')
        setting = data.get('setting', '')
        page_number = data.get('page_number', 1)
        image_api = data.get('image_api', 'dalle')
        
        # Dump the entire data for debugging
        logger.info(f"generate_page_image - FULL DATA: {data}")
        
        # Karakter özelliklerini al
        character_description = ""
        character_age = data.get('character_age', '')
        character_gender = data.get('character_gender', '')
        character_hair_type = data.get('character_hair_type', '')
        character_hair_color = data.get('character_hair_color', '')
        character_skin_color = data.get('character_skin_color', '')
        
        # Debug için tüm özellikleri logla
        logger.info(f"generate_page_image - character_age: {character_age}")
        logger.info(f"generate_page_image - character_gender: {character_gender}")
        logger.info(f"generate_page_image - character_hair_type: {character_hair_type}")
        logger.info(f"generate_page_image - character_hair_color: {character_hair_color}")
        logger.info(f"generate_page_image - character_skin_color: {character_skin_color}")
        
        if character_age:
            character_description += f"{character_age} yaşında, "
        
        if character_gender:
            character_description += f"{character_gender}, "
        
        hair_parts = []
        if character_hair_type:
            hair_parts.append(character_hair_type)
        if character_hair_color:
            hair_parts.append(character_hair_color)
        
        if hair_parts:
            character_description += f"{' '.join(hair_parts)} saçlı, "
        
        if character_skin_color:
            character_description += f"{character_skin_color} tenli, "
        
        # Son virgülü ve boşluğu kaldır
        if character_description:
            character_description = character_description.rstrip(", ")
            character_description = f"({character_description}) "
        
        logger.info(f"Sayfa {page_number} için görsel isteği alındı")
        
        # Görsel oluşturma promptu hazırla
        image_prompt = f"{character_name} adlı {character_description}{character_type} karakteri {setting} ortamında: {page_text}"
        logger.info(f"Oluşturulan prompt: {image_prompt[:100]}...")
        # Prompt'u tamamen logla
        prompt_logger.info(f"Page {page_number} image prompt: {image_prompt}")
        
        # Görseli oluştur
        image_data = generate_image_for_section(image_prompt, image_api)
        
        if not image_data:
            logger.warning(f"Sayfa {page_number} için görsel oluşturulamadı")
            return jsonify({"error": "Görsel oluşturulamadı"}), 500
        
        logger.info(f"Sayfa {page_number} için görsel başarıyla oluşturuldu")
        
        # Base64 verisi olarak dön
        return jsonify({
            "image_data": image_data,
            "image_url": f"data:image/jpeg;base64,{image_data}"
        })
        
    except Exception as e:
        logger.error(f"Sayfa görseli oluşturma hatası: {str(e)}")
        logger.error(traceback.format_exc())
        return jsonify({'error': str(e)}), 500

def create_word_document(tale_text, images):
    """Masal metni ve görsellerden Word dosyası oluşturur"""
    try:
        # Word dosyası oluştur
        doc = Document()
        
        # Başlık ekle
        doc.add_heading('Masal Dünyası', 0)
        
        # Masal metnini ve görselleri ekle
        paragraphs = tale_text.split('\n\n')
        
        # Eğer görseller ve paragraflar varsa, her ikisini birleştir
        if images and paragraphs:
            # Sayfaları ve görselleri yan yana ekle
            for i, paragraph_text in enumerate(paragraphs):
                # Metni ekle
                p = doc.add_paragraph(paragraph_text)
                
                # Sayfa sonuna uygun görsel varsa ekle
                if i < len(images):
                    try:
                        # Base64 görselini dosyaya dönüştür
                        image_data = images[i]
                        logger.info(f"Görsel {i+1} işleniyor - Veri uzunluğu: {len(image_data[:20])}...")
                        
                        image_bytes = base64.b64decode(image_data)
                        
                        # Geçici dosya oluştur
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                            temp_path = temp_file.name
                            temp_file.write(image_bytes)
                        
                        # Görseli ekle
                        try:
                            doc.add_picture(temp_path, width=Inches(6))
                            logger.info(f"Görsel {i+1} Word belgesine eklendi")
                        except Exception as img_error:
                            logger.error(f"Görsel {i+1} eklenirken hata: {str(img_error)}")
                        
                        # Geçici dosyayı sil
                        os.unlink(temp_path)
                    except Exception as img_e:
                        logger.error(f"Görsel {i+1} işlenirken hata: {str(img_e)}")
                
                # Sayfa sonu ekle (son sayfa hariç)
                if i < len(paragraphs) - 1:
                    doc.add_page_break()
        else:
            # Sadece metni ekle (görsel yoksa)
            doc.add_paragraph(tale_text)
            logger.info("Görsel olmadığı için sadece metin eklendi")
        
        # Word dosyasını kaydet
        doc_path = tempfile.mktemp(suffix='.docx')
        doc.save(doc_path)
        
        logger.info(f"Word belgesi oluşturuldu: {doc_path}, Boyut: {os.path.getsize(doc_path)} bytes")
        return doc_path
    except Exception as e:
        logger.error(f"Word dosyası oluşturma hatası: {e}")
        logger.error(traceback.format_exc())
        raise

def create_audio(text):
    """Metinden ses dosyası oluşturur - basitleştirilmiş, pydub gerektirmeyen sürüm"""
    try:
        # Geçici ses dosyası oluştur
        with tempfile.NamedTemporaryFile(suffix='.mp3', delete=False) as f:
            audio_path = f.name
        
        # Metin içeriğinin uzunluğunu logla
        text_words = len(text.split())
        logger.info(f"Ses oluşturulacak metin: {text_words} kelime, {len(text)} karakter")
        logger.debug(f"Metin başlangıcı: {text[:min(50, len(text))]}...")
        
        # Metni sese dönüştür - uzun metinler için gTTS kendi içinde parçalama yapıyor
        tts = gTTS(text=text, lang='tr', slow=False)
        
        # Dosyaya kaydet
        tts.save(audio_path)
        
        # Dosya boyutunu kontrol et
        file_size = os.path.getsize(audio_path)
        logger.info(f"Ses dosyası oluşturuldu: {audio_path} (Boyut: {file_size} bytes)")
        
        return audio_path
    except Exception as e:
        logger.error(f"Ses oluşturma hatası: {e}")
        logger.error(traceback.format_exc())
        raise

def generate_tale_text(character_name, character_type, setting, theme, word_limit, text_api='openai', character_attributes=None):
    """Gemini API veya OpenAI API kullanarak masal metni oluşturur"""
    try:
        # Karakter bilgilerini al
        character_description = ""
        if character_attributes:
            description_parts = []
            age = character_attributes.get('character_age')
            gender = character_attributes.get('character_gender')
            hair_color = character_attributes.get('character_hair_color')
            hair_type = character_attributes.get('character_hair_type')
            skin_color = character_attributes.get('character_skin_color')
            
            if age:
                description_parts.append(f"{age} yaşında")
            if gender:
                description_parts.append(f"{gender}")
            if hair_color and hair_type:
                description_parts.append(f"{hair_type} {hair_color} saçlı")
            elif hair_color:
                description_parts.append(f"{hair_color} saçlı")
            elif hair_type:
                description_parts.append(f"{hair_type} saçlı")
            if skin_color:
                description_parts.append(f"{skin_color} tenli")
            
            if description_parts:
                character_description = f"Fiziksel özellikler: {', '.join(description_parts)}. "
                logger.info(f"Karakter detayları: {character_description}")
        
        # Kelime limitini 1.4 ile çarp - daha gerçekçi sayılara ulaşmak için
        adjusted_word_limit = int(word_limit * 1.4)
        logger.info(f"Orijinal kelime limiti: {word_limit}, Ayarlanmış limit: {adjusted_word_limit}")
        
        # Kullanıcı seçimine göre API kullan
        if text_api == 'openai' and openai_client:
            try:
                return generate_tale_text_with_openai(character_name, character_type, setting, theme, adjusted_word_limit, character_description)
            except Exception as e:
                logger.warning(f"OpenAI ile masal metni oluşturulamadı, Gemini denenecek: {str(e)}")
                # OpenAI ile başarısız olursa Gemini'ye devam et
        
        # Gemini API ile devam et
        logger.info("Gemini API ile masal metni oluşturuluyor...")
        
        # Desteklenen modelleri kullan
        logger.info("Gemini modeli kullanılıyor...")
        
        # İlk tercih edilen model
        try:
            model = genai.GenerativeModel("models/gemini-2.0-flash-001")
            logger.info("models/gemini-2.0-flash-001 modeli başarıyla yüklendi")
        except Exception as e1:
            logger.warning(f"models/gemini-2.0-flash-001 modeli yüklenemedi: {str(e1)}")
            
            # İkinci tercih edilen model
            try:
                model = genai.GenerativeModel("models/gemini-2.0-flash-lite-001")
                logger.info("models/gemini-2.0-flash-lite-001 modeli başarıyla yüklendi")
            except Exception as e2:
                logger.warning(f"models/gemini-2.0-flash-lite-001 modeli yüklenemedi: {str(e2)}")
                
                # Son çare olarak bilinen çalışan modeli dene
                try:
                    model = genai.GenerativeModel("models/gemini-1.5-pro")
                    logger.info("models/gemini-1.5-pro modeli başarıyla yüklendi")
                except Exception as e3:
                    logger.error(f"Hiçbir Gemini modeli yüklenemedi: {str(e3)}")
                    if not openai_client:
                        raise ValueError("Hem Gemini hem de OpenAI API kullanılamıyor. Lütfen API anahtarlarını kontrol edin.")
                    else:
                        # Buraya düşerse, OpenAI zaten denenmiş ve başarısız olmuş demektir
                        raise ValueError("Metin oluşturma için hiçbir API kullanılamıyor.")
        
        prompt = f"""
        GÖREV: Tam olarak {adjusted_word_limit} kelimeden oluşan bir çocuk masalı yaz.

        ÖNEMLİ TALİMATLAR:
        1. Hikaye TAM OLARAK {adjusted_word_limit} kelime içermeli
        2. Hikaye şunları içermeli:
           - Ana karakter: {character_name} adında bir {character_type}
           - {character_description}
           - Ortam: {setting}
           - Tema: {theme}
        3. Çocuk dostu ve eğitici olmalı (7-10 yaş)
        4. Basit Türkçe kullan
        5. Kelime sayımını üç kez kontrol et
        6. Başlık EKLEME
        7. Ne bir kelime fazla, ne bir kelime eksik olmalı

        ÇOK ÖNEMLİ:
        - Kelime sayısını metnin kendisinde belirtme (yani metinde "Bu hikaye {adjusted_word_limit} kelimedir" gibi ifadeler kullanma)
        - Asla "Unutmayın bu masal tam olarak X kelime içeriyor" veya benzeri ifadeler ekleme
        - Sadece masal içeriğini yaz, başka açıklama ekleme
        """
        
        response = model.generate_content(prompt)
        tale_text = response.text.strip()
        
        # Kelime sayısı kontrolü
        words = tale_text.split()
        word_count = len(words)
        logger.info(f"Model tarafından üretilen kelime sayısı: {word_count}")
        
        if word_count > adjusted_word_limit * 1.2:  # %20 tolerans
            logger.info(f"Kelime sayısı fazla, {adjusted_word_limit} kelimeye kısaltılıyor...")
            tale_text = ' '.join(words[:adjusted_word_limit])
        elif word_count < adjusted_word_limit * 0.8:  # Kelime sayısı %20'den fazla az ise
            logger.warning(f"Kelime sayısı çok az ({word_count}), istenen: {adjusted_word_limit}. Yeniden deneniyor...")
            
            # Daha sıkı kontrollerle bir retry prompt oluşturalım
            retry_prompt = f"""
            GÖREV: Tam olarak {adjusted_word_limit} kelimeden oluşan bir çocuk masalı yaz.

            ÖNEMLİ TALİMATLAR (DİKKATLİCE UYGULANACAK):
            1. Hikaye TAM OLARAK {adjusted_word_limit} kelime içermeli - kelime sayacı kullanarak ÜÇ KEZ sayımı doğrula
            2. Hikaye şunları içermeli:
               - Ana karakter: {character_name} adında bir {character_type}
               - {character_description}
               - Ortam: {setting}
               - Tema: {theme}
            3. Çocuk dostu ve eğitici olmalı (7-10 yaş)
            4. Basit Türkçe kullan
            5. Başlık EKLEME
            6. Kelime sayısını metnin içinde belirtme (yani "bu masal X kelimedir" gibi cümleler kullanma)
            7. Masal {adjusted_word_limit} KELİMEDEN NE BİR FAZLA NE BİR EKSİK olmalı
            8. Sadece masal metnini ver, açıklama veya not ekleme
            9. Asla "Unutmayın bu masal tam olarak X kelime içeriyor" veya benzeri ifadeler ekleme
            """
            
            try:
                logger.info("Gemini ile yeniden deneme yapılıyor...")
                retry_response = model.generate_content(retry_prompt)
                retry_text = retry_response.text.strip()
                
                # Yeniden deneme sonucunu kontrol et
                retry_word_count = len(retry_text.split())
                logger.info(f"Yeniden deneme sonucu kelime sayısı: {retry_word_count}")
                
                if abs(retry_word_count - adjusted_word_limit) < abs(word_count - adjusted_word_limit):
                    logger.info(f"Yeniden deneme daha iyi sonuç verdi. Önceki: {word_count}, Yeni: {retry_word_count}, Hedef: {adjusted_word_limit}")
                    return retry_text
                else:
                    logger.info(f"Yeniden deneme daha iyi sonuç vermedi. Önceki: {word_count}, Yeni: {retry_word_count}, Hedef: {adjusted_word_limit}")
            except Exception as retry_error:
                logger.error(f"Yeniden deneme sırasında hata: {str(retry_error)}")
                logger.debug(f"Önceki sonuç kullanılıyor (kelime sayısı: {word_count})")
        
        return tale_text
    
    except Exception as e:
        logger.error(f"Masal metni oluşturulurken hata: {str(e)}")
        logger.error(traceback.format_exc())
        return f"Masal oluşturulamadı. Hata: {str(e)}"
        
def generate_tale_text_with_openai(character_name, character_type, setting, theme, word_limit, character_description=""):
    """OpenAI API kullanarak masal metni oluşturur"""
    logger.info("OpenAI API ile masal metni oluşturuluyor...")
    
    prompt = f"""
    GÖREV: Tam olarak {word_limit} kelimeden oluşan bir çocuk masalı yaz.

    ÖNEMLİ TALİMATLAR:
    1. Hikaye TAM OLARAK {word_limit} kelime içermeli
    2. Hikaye şunları içermeli:
       - Ana karakter: {character_name} adında bir {character_type}
       - {character_description}
       - Ortam: {setting}
       - Tema: {theme}
    3. Çocuk dostu ve eğitici olmalı (7-10 yaş)
    4. Basit Türkçe kullan
    5. Kelime sayımını üç kez kontrol et
    6. Başlık EKLEME
    7. Ne bir kelime fazla, ne bir kelime eksik olmalı

    ÇOK ÖNEMLİ:
    - Kelime sayısını metnin kendisinde belirtme (metinde "Bu hikaye {word_limit} kelimedir" gibi ifadeler kullanma)
    - Asla "Unutmayın bu masal tam olarak X kelime içeriyor" veya benzeri ifadeler ekleme
    - Sadece masal içeriğini yaz, başka açıklama ekleme
    """
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini-2024-07-18",  # Daha yeni ve gelişmiş bir model kullan
            messages=[
                {"role": "system", "content": "Sen çocuklar için masal yazan bir yazarsın. Eğitici, eğlenceli ve çocuk dostu masallar yazarsın."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=min(4000, word_limit * 10),  # Kelime başına daha fazla token verelim ama limit koyalım
            temperature=0.7,
            presence_penalty=0.1,  # Tekrarları önlemek için hafif bir presence penalty ekleyelim
            frequency_penalty=0.1  # Tekrarları önlemek için hafif bir frequency penalty ekleyelim
        )
        
        tale_text = response.choices[0].message.content.strip()
        
        # Kelime sayısı kontrolü
        words = tale_text.split()
        word_count = len(words)
        logger.info(f"OpenAI tarafından üretilen kelime sayısı: {word_count}")
        
        if word_count > word_limit * 1.2:  # %20 tolerans
            logger.info(f"Kelime sayısı fazla, {word_limit} kelimeye kısaltılıyor...")
            tale_text = ' '.join(words[:word_limit])
        elif word_count < word_limit * 0.8:  # Kelime sayısı %20'den fazla az ise
            logger.warning(f"Kelime sayısı çok az ({word_count}), istenen: {word_limit}. Yeniden deneniyor...")
            
            # Yeni bir prompt oluştur ve daha net şekilde kelime sayısını vurgula
            retry_prompt = f"""
            GÖREV: Tam olarak {word_limit} kelimeden oluşan bir çocuk masalı yaz.

            ÖNEMLİ TALİMATLAR (DİKKATLİCE UYGULANACAK):
            1. Hikaye TAM OLARAK {word_limit} kelime içermeli - kelime sayacı kullanarak ÜÇ KEZ sayımı doğrula
            2. Hikaye şunları içermeli:
               - Ana karakter: {character_name} adında bir {character_type}
               - Ortam: {setting}
               - Tema: {theme}
            3. Çocuk dostu ve eğitici olmalı (7-10 yaş)
            4. Basit Türkçe kullan
            5. Başlık EKLEME
            6. Kelime sayısını metnin içinde belirtme (yani "bu masal X kelimedir" gibi cümleler kullanma)
            7. Masal {word_limit} KELİMEDEN NE BİR FAZLA NE BİR EKSİK olmalı
            8. "Unutmayın bu masal tam olarak X kelime içeriyor" gibi ifadeler asla kullanma
            
            ÇOK ÖNEMLİ: Sadece masal metnini gönder. Başlık, açıklama veya kelime sayısı bildirimi ekleme.
            """
            
            try:
                # Yeniden deneme - Daha iyi bir model kullan
                retry_response = openai_client.chat.completions.create(
                    model="gpt-4-turbo-2024-04-09",  # İlk deneme başarısız olursa daha güçlü modele geç
                    messages=[
                        {"role": "system", "content": "Sen kelime sayısı limitlerini tam olarak izleyen bir yazarsın. Verilen kelime sayısı limitlerini daima tam olarak uygularsın."},
                        {"role": "user", "content": retry_prompt}
                    ],
                    max_tokens=min(4000, word_limit * 10),  # Daha fazla token, ama limite uygun
                    temperature=0.5,  # Daha az yaratıcılık (daha kurallara uygun)
                    presence_penalty=0.2,
                    frequency_penalty=0.2
                )
                
                # Yeni yanıtı kontrol et
                retry_text = retry_response.choices[0].message.content.strip()
                retry_words = retry_text.split()
                retry_count = len(retry_words)
                
                logger.info(f"Yeniden deneme sonucu kelime sayısı: {retry_count}")
                
                # Eğer yeni deneme daha iyi sonuç verdiyse, onu kullan
                if abs(retry_count - word_limit) < abs(word_count - word_limit):
                    logger.info(f"Yeniden deneme daha iyi sonuç verdi, bu metin kullanılacak")
                    tale_text = retry_text
                    # Kelime sınırı kontrolü hala gerekli olabilir
                    if retry_count > word_limit * 1.2:
                        tale_text = ' '.join(retry_words[:word_limit])
                else:
                    logger.info(f"İlk deneme daha iyi sonuç verdi, orijinal metin kullanılacak")
            except Exception as retry_e:
                logger.error(f"Yeniden deneme sırasında hata: {str(retry_e)}")
                # Hata durumunda orijinal metni kullan
        
        return tale_text
        
    except Exception as e:
        logger.error(f"OpenAI ile masal metni oluşturulurken hata: {str(e)}")
        logger.error(traceback.format_exc())
        raise

def split_text_into_sections(text, words_per_section):
    """Metni belirli kelime sayısına göre bölümlere ayırır"""
    words = text.split()
    sections = []
    
    for i in range(0, len(words), words_per_section):
        section = ' '.join(words[i:i + words_per_section])
        sections.append(section)
    
    return sections

def generate_image_for_section(section_text, image_api='dalle'):
    """Bölüm metni için görsel oluşturur"""
    try:
        # Sadece DALL-E kullan, OpenAI API anahtarı kontrol et
        if openai_client:
            logger.info("DALL-E API kullanılıyor...")
            return generate_image_with_dalle(section_text)
        else:
            logger.warning("OpenAI API anahtarı yok, placeholder görüntü oluşturuluyor...")
            return create_placeholder_image(section_text)
    except Exception as e:
        logger.error(f"Görsel oluşturma hatası: {e}")
        logger.error(traceback.format_exc())
        # Hata durumunda placeholder görüntü oluştur
        return create_placeholder_image(section_text)

def generate_image_with_dalle(prompt):
    """OpenAI DALL-E API kullanarak görsel oluşturur"""
    try:
        # Prompt'u çocuk dostu hale getir ve yazı içermemesini sağla
        enhanced_prompt = f"Turkish children's book style illustration with NO TEXT. Create a colorful, cartoon-style image with ABSOLUTELY NO words, text, or speech bubbles. The image must NOT contain any letters, alphabet characters, or written text: {prompt}"
        logger.info(f"DALL-E prompt: {enhanced_prompt[:100]}...")
        # Prompt'u tamamen logla
        prompt_logger.info(f"DALL-E prompt: {enhanced_prompt}")
        
        # Rate limit kontrolü için basit hız sınırlama
        # En son DALL-E API çağrısı zamanını kontrol et
        current_time = time.time()
        if hasattr(generate_image_with_dalle, 'last_call_time'):
            time_since_last_call = current_time - generate_image_with_dalle.last_call_time
            # DALL-E'nin rate limiti dakikada 5 istek olduğu için en az 12 saniye bekleyelim
            if time_since_last_call < 12:  # Saniye cinsinden bekleme süresi
                wait_time = 12 - time_since_last_call
                logger.info(f"DALL-E rate limit koruması: {wait_time:.2f} saniye bekleniyor...")
                time.sleep(wait_time)
        
        # Bu çağrı zamanını kaydet
        generate_image_with_dalle.last_call_time = time.time()
        
        response = openai_client.images.generate(
            model="dall-e-3",
            prompt=enhanced_prompt,
            size="1024x1024",
            quality="standard",
            n=1,
            style="vivid"  # "vivid" veya "natural" olabilir
        )
        image_url = response.data[0].url
        logger.info(f"DALL-E görsel URL'si oluşturuldu: {image_url[:50]}...")
        
        # Resmi indir
        image_response = requests.get(image_url)
        image_data = image_response.content
        
        # Base64 formatına dönüştür
        base64_image = base64.b64encode(image_data).decode('utf-8')
        logger.info("DALL-E görsel başarıyla oluşturuldu ve Base64 formatına dönüştürüldü.")
        return base64_image
    
    except Exception as e:
        logger.error(f"OpenAI ile görsel oluşturma hatası: {e}")
        logger.error(traceback.format_exc())
        
        # Farklı hata türlerine göre işlem yap
        if isinstance(e, openai.RateLimitError):
            try:
                # Rate limit hatası - daha uzun süre bekle ve yeniden dene
                wait_time = 30  # Rate limit için 30 saniye bekle 
                logger.info(f"DALL-E rate limit hatası, {wait_time} saniye bekleniyor ve yeniden deneniyor...")
                time.sleep(wait_time)
                
                # Yeniden dene
                generate_image_with_dalle.last_call_time = time.time()  # Son çağrı zamanını güncelle
                
                response = openai_client.images.generate(
                    model="dall-e-3",
                    prompt=enhanced_prompt,
                    size="1024x1024",
                    quality="standard",
                    n=1,
                    style="vivid"
                )
                image_url = response.data[0].url
                logger.info(f"DALL-E görsel URL'si yeniden deneme sonrası oluşturuldu: {image_url[:50]}...")
                
                # Resmi indir
                image_response = requests.get(image_url)
                image_data = image_response.content
                
                # Base64 formatına dönüştür
                base64_image = base64.b64encode(image_data).decode('utf-8')
                logger.info("DALL-E görsel başarıyla oluşturuldu ve Base64 formatına dönüştürüldü (yeniden deneme sonrası).")
                return base64_image
                
            except Exception as retry_error:
                logger.error(f"DALL-E yeniden deneme hatası: {retry_error}")
                logger.error(traceback.format_exc())
        
        elif isinstance(e, openai.BadRequestError):
            # İçerik politikası hatası veya diğer Bad Request hataları
            logger.warning(f"DALL-E içerik politikası hatası veya geçersiz istek: {str(e)}")
            logger.info("İçerik politikası nedeniyle placeholder görüntü oluşturuluyor...")
            
            # İçerik verisi çok özel olabilir, daha genel bir prompt deneyelim
            try:
                # Daha genel ve güvenli bir prompt oluştur
                safe_prompt = "Çocuk dostu, renkli, çizgi film tarzında: Güzel bir orman manzarası, ağaçlar ve çiçekler"
                
                # Saniye kadar bekle - rate limit için
                time.sleep(15)
                
                # Yeniden dene
                generate_image_with_dalle.last_call_time = time.time()
                
                response = openai_client.images.generate(
                    model="dall-e-3",
                    prompt=safe_prompt,
                    size="1024x1024",
                    quality="standard",
                    n=1,
                    style="vivid"
                )
                image_url = response.data[0].url
                logger.info(f"DALL-E güvenli prompt ile görsel URL'si oluşturuldu: {image_url[:50]}...")
                
                # Resmi indir
                image_response = requests.get(image_url)
                image_data = image_response.content
                
                # Base64 formatına dönüştür
                base64_image = base64.b64encode(image_data).decode('utf-8')
                logger.info("DALL-E görsel güvenli prompt ile oluşturuldu.")
                return base64_image
                
            except Exception as safe_retry_error:
                logger.error(f"Güvenli prompt ile yeniden deneme hatası: {safe_retry_error}")
        
        # Hata durumunda placeholder görüntü oluştur
        logger.info("DALL-E hatası nedeniyle placeholder görüntü oluşturuluyor...")
        return create_placeholder_image(prompt)

def generate_image_with_gemini(section_text):
    """Gemini API kullanarak görsel oluşturur"""
    try:
        # İlk tercih edilen model
        try:
            model = genai.GenerativeModel("models/gemini-2.0-flash-001")
            logger.info("Görsel için models/gemini-2.0-flash-001 modeli kullanılıyor")
        except Exception as e1:
            logger.warning(f"Görsel için models/gemini-2.0-flash-001 modeli yüklenemedi: {str(e1)}")
            
            # İkinci tercih edilen model
            try:
                model = genai.GenerativeModel("models/gemini-2.0-flash-lite-001")
                logger.info("Görsel için models/gemini-2.0-flash-lite-001 modeli kullanılıyor")
            except Exception as e2:
                logger.warning(f"Görsel için models/gemini-2.0-flash-lite-001 modeli yüklenemedi: {str(e2)}")
                
                # Son çare olarak bilinen çalışan modeli dene
                model = genai.GenerativeModel("models/gemini-1.5-pro")
                logger.info("Görsel için models/gemini-1.5-pro modeli kullanılıyor")
        
        prompt = f"""
        Lütfen aşağıdaki metne uygun, çocuk kitabı tarzında, renkli ve sevimli bir illüstrasyon oluştur:
        
        {section_text}
        
        İllüstrasyon 3-7 yaş arası çocuklar için uygun, canlı renkli ve detaylı olmalı.
        Tarz olarak Disney/Pixar animasyon filmlerine benzer, sevimli karakterler içermeli.
        Görsel yüksek çözünürlüklü ve net olmalı.
        """
        
        response = model.generate_content(prompt)
        
        # Yanıtı kontrol et ve görüntü verisini çıkar
        for candidate in response.candidates:
            for part in candidate.content.parts:
                if hasattr(part, 'inline_data') and part.inline_data:
                    # Base64 kodlu görüntüyü döndür
                    logger.info("Gemini görsel başarıyla oluşturuldu.")
                    return part.inline_data.data
        
        # Eğer görüntü oluşturulamazsa, boş bir görüntü oluştur
        logger.warning("Gemini görsel oluşturamadı, placeholder görüntü oluşturuluyor.")
        return create_placeholder_image(section_text)
    
    except Exception as e:
        logger.error(f"Görsel oluşturma hatası: {e}")
        logger.error(traceback.format_exc())
        # Hata durumunda da placeholder görüntü oluştur
        return create_placeholder_image(section_text)

def create_placeholder_image(text):
    """Metinden daha çekici bir placeholder görüntü oluşturur"""
    try:
        # Boyutları ve arkaplan rengini belirle - açık mavi
        width, height = 1024, 1024
        image = Image.new('RGB', (width, height), color=(173, 216, 230))
        
        # Görsel iyileştirmeleri ekle - basit bir gökkuşağı arka plan
        from PIL import ImageDraw, ImageFont
        
        draw = ImageDraw.Draw(image)
        
        # Renkli çerçeve çiz
        border_width = 20
        draw.rectangle(
            [(border_width, border_width), (width - border_width, height - border_width)],
            outline=(255, 255, 255),
            width=5
        )
        
        # Gökkuşağı renkli süsleme çiz
        colors = [
            (255, 105, 180),  # Pembe
            (255, 165, 0),    # Turuncu
            (255, 215, 0),    # Altın sarısı
            (144, 238, 144),  # Açık yeşil
            (135, 206, 250)   # Açık mavi
        ]
        
        # Renkli daireler çiz
        for i, color in enumerate(colors):
            pos_x = width // 4 + (i * 50) % (width // 2)
            pos_y = height // 4 + (i * 50) % (height // 2)
            circle_size = 30 + (i * 5)
            draw.ellipse(
                [(pos_x, pos_y), (pos_x + circle_size, pos_y + circle_size)],
                fill=color,
                outline=(255, 255, 255)
            )
        
        # Metin ekle
        try:
            # Arial font kullan, yoksa varsayılan font
            font = ImageFont.truetype("Arial", 24)
        except IOError:
            font = ImageFont.load_default()
            
        text_to_display = "Resim yüklenemedi. Yeni bir görsel oluşturuluyor..."
        
        # Metin için arka plan dikdörtgeni çiz
        text_bbox = draw.textbbox((0, 0), text_to_display, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        text_pos_x = (width - text_width) // 2
        text_pos_y = height // 2 - text_height // 2
        
        # Metin arkaplanı
        padding = 10
        draw.rectangle(
            [
                (text_pos_x - padding, text_pos_y - padding),
                (text_pos_x + text_width + padding, text_pos_y + text_height + padding)
            ],
            fill=(255, 255, 255, 128)
        )
        
        # Metni ortaya yerleştir
        draw.text(
            (text_pos_x, text_pos_y),
            text_to_display,
            font=font,
            fill=(0, 0, 0)
        )
        
        # Görüntüyü base64 formatına dönüştür
        buffer = io.BytesIO()
        image.save(buffer, format="PNG")
        base64_image = base64.b64encode(buffer.getvalue()).decode('utf-8')
        
        logger.info("Geliştirilmiş placeholder görüntü başarıyla oluşturuldu.")
        return base64_image
    except Exception as e:
        logger.error(f"Placeholder görüntü oluşturma hatası: {e}")
        logger.error(traceback.format_exc())
        
        # Hata durumunda çok basit bir görüntü oluştur
        try:
            simple_image = Image.new('RGB', (1024, 1024), color=(240, 248, 255))
            buffer = io.BytesIO()
            simple_image.save(buffer, format="PNG")
            base64_image = base64.b64encode(buffer.getvalue()).decode('utf-8')
            return base64_image
        except:
            return None

# Ses efekti fonksiyonu kaldırıldı

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8500, debug=True)
